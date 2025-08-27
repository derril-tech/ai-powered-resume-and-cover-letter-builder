from typing import Dict, Any
from docx import Document


class DocxRenderer:
    async def render(self, content: Dict[str, Any]) -> bytes:
        doc = Document()
        title = content.get("title") or content.get("name") or "Resume"
        doc.add_heading(title, 0)

        for section, value in content.items():
            if section == "title" or section == "name":
                continue
            doc.add_heading(section.title(), level=1)
            if isinstance(value, str):
                doc.add_paragraph(value)
            elif isinstance(value, list):
                for item in value:
                    doc.add_paragraph(str(item), style="List Bullet")
            elif isinstance(value, dict):
                for k, v in value.items():
                    doc.add_paragraph(f"{k}: {v}")

        from io import BytesIO

        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()


