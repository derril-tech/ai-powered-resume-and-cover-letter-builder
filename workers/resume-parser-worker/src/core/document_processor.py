"""
Document Processor
Handles extraction of text from various document formats (PDF, DOCX, TXT).
"""

import logging
import re
from typing import Optional, Dict, Any
import io
import base64

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Processes documents and extracts text content"""

    def __init__(self):
        self.supported_formats = {
            'application/pdf': 'pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
            'application/msword': 'doc',
            'text/plain': 'txt',
        }

    async def extract_text(
        self,
        file_content: bytes,
        content_type: str,
        filename: Optional[str] = None,
    ) -> str:
        """
        Extract text from document based on content type

        Args:
            file_content: Raw file bytes
            content_type: MIME type of the file
            filename: Original filename (optional, for format detection)

        Returns:
            Extracted text content
        """
        if content_type not in self.supported_formats:
            raise ValueError(f"Unsupported content type: {content_type}")

        format_type = self.supported_formats[content_type]

        try:
            if format_type == 'pdf':
                return await self._extract_pdf_text(file_content)
            elif format_type == 'docx':
                return await self._extract_docx_text(file_content)
            elif format_type == 'doc':
                return await self._extract_doc_text(file_content)
            elif format_type == 'txt':
                return await self._extract_txt_text(file_content)
            else:
                raise ValueError(f"No handler for format: {format_type}")

        except Exception as e:
            logger.error(f"Failed to extract text from {format_type} document", error=str(e))
            # Try fallback extraction
            return await self._fallback_text_extraction(file_content)

    async def _extract_pdf_text(self, file_content: bytes) -> str:
        """Extract text from PDF files"""
        try:
            import PyPDF2
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))

            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"

            return self._clean_extracted_text(text)

        except ImportError:
            logger.warning("PyPDF2 not available, trying pdfminer")
            return await self._extract_pdf_with_pdfminer(file_content)
        except Exception as e:
            logger.warning(f"PyPDF2 extraction failed: {e}")
            return await self._extract_pdf_with_pdfminer(file_content)

    async def _extract_pdf_with_pdfminer(self, file_content: bytes) -> str:
        """Fallback PDF extraction using pdfminer"""
        try:
            from pdfminer.high_level import extract_text
            text = extract_text(io.BytesIO(file_content))
            return self._clean_extracted_text(text)
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            raise

    async def _extract_docx_text(self, file_content: bytes) -> str:
        """Extract text from DOCX files"""
        try:
            from docx import Document
            doc = Document(io.BytesIO(file_content))

            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            return self._clean_extracted_text(text)

        except ImportError:
            logger.warning("python-docx not available, trying docx2txt")
            return await self._extract_docx_with_docx2txt(file_content)
        except Exception as e:
            logger.warning(f"DOCX extraction failed: {e}")
            return await self._extract_docx_with_docx2txt(file_content)

    async def _extract_docx_with_docx2txt(self, file_content: bytes) -> str:
        """Fallback DOCX extraction using docx2txt"""
        try:
            import docx2txt
            text = docx2txt.process(io.BytesIO(file_content))
            return self._clean_extracted_text(text)
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise

    async def _extract_doc_text(self, file_content: bytes) -> str:
        """Extract text from DOC files"""
        try:
            import textract
            text = textract.process(io.BytesIO(file_content))
            return self._clean_extracted_text(text.decode('utf-8'))
        except Exception as e:
            logger.error(f"DOC extraction failed: {e}")
            raise

    async def _extract_txt_text(self, file_content: bytes) -> str:
        """Extract text from TXT files"""
        try:
            text = file_content.decode('utf-8')
            return self._clean_extracted_text(text)
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    text = file_content.decode(encoding)
                    return self._clean_extracted_text(text)
                except UnicodeDecodeError:
                    continue

            raise ValueError("Could not decode text file with any supported encoding")

    async def _fallback_text_extraction(self, file_content: bytes) -> str:
        """Fallback text extraction using basic methods"""
        try:
            # Try to decode as UTF-8
            text = file_content.decode('utf-8')
            return self._clean_extracted_text(text)
        except UnicodeDecodeError:
            # Try to extract readable strings
            text = ""
            current_string = ""

            for byte in file_content:
                if 32 <= byte <= 126:  # Printable ASCII characters
                    current_string += chr(byte)
                else:
                    if len(current_string) >= 4:  # Only keep strings of reasonable length
                        text += current_string + " "
                    current_string = ""

            return self._clean_extracted_text(text)

    def _clean_extracted_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""

        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)

        # Remove non-printable characters
        text = re.sub(r'[^\x20-\x7E\n\r\t]', '', text)

        # Fix common OCR issues
        text = re.sub(r'([a-zA-Z])\s*([.,!?])', r'\1\2', text)  # Remove spaces before punctuation
        text = re.sub(r'([.,!?])\s*([a-zA-Z])', r'\1 \2', text)  # Add spaces after punctuation

        # Remove excessive line breaks
        text = re.sub(r'\n{3,}', '\n\n', text)

        # Fix common spacing issues
        text = re.sub(r'\s*-\s*', '-', text)  # Fix hyphens
        text = re.sub(r'\s*/\s*', '/', text)  # Fix slashes

        return text.strip()

    async def detect_document_type(self, filename: Optional[str], content_type: str) -> str:
        """Detect document type from filename and content type"""
        if filename:
            filename_lower = filename.lower()
            if filename_lower.endswith('.pdf'):
                return 'pdf'
            elif filename_lower.endswith('.docx'):
                return 'docx'
            elif filename_lower.endswith('.doc'):
                return 'doc'
            elif filename_lower.endswith('.txt'):
                return 'txt'

        # Fallback to content type
        return self.supported_formats.get(content_type, 'unknown')

    async def get_document_info(self, file_content: bytes, filename: Optional[str], content_type: str) -> Dict[str, Any]:
        """Get document metadata"""
        info = {
            'filename': filename,
            'content_type': content_type,
            'size_bytes': len(file_content),
            'detected_type': await self.detect_document_type(filename, content_type),
        }

        # Try to extract basic document info
        if info['detected_type'] == 'pdf':
            info.update(await self._get_pdf_info(file_content))
        elif info['detected_type'] in ['docx', 'doc']:
            info.update(await self._get_word_info(file_content))

        return info

    async def _get_pdf_info(self, file_content: bytes) -> Dict[str, Any]:
        """Get PDF-specific information"""
        try:
            import PyPDF2
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))

            return {
                'page_count': len(pdf_reader.pages),
                'has_metadata': pdf_reader.metadata is not None,
                'metadata': dict(pdf_reader.metadata) if pdf_reader.metadata else None,
            }
        except Exception:
            return {}

    async def _get_word_info(self, file_content: bytes) -> Dict[str, Any]:
        """Get Word document information"""
        try:
            from docx import Document
            doc = Document(io.BytesIO(file_content))

            return {
                'paragraph_count': len(doc.paragraphs),
                'has_tables': len(doc.tables) > 0,
                'table_count': len(doc.tables),
            }
        except Exception:
            return {}

    async def validate_document(self, file_content: bytes, content_type: str) -> Dict[str, Any]:
        """Validate document and return validation results"""
        validation = {
            'is_valid': True,
            'issues': [],
            'warnings': [],
            'recommendations': [],
        }

        # Check file size
        size_mb = len(file_content) / (1024 * 1024)
        if size_mb > 10:
            validation['issues'].append('File size exceeds 10MB limit')
            validation['is_valid'] = False
        elif size_mb > 5:
            validation['warnings'].append('Large file size may affect processing speed')

        # Check if file appears to be text-based
        try:
            text = await self.extract_text(file_content, content_type)
            if len(text.strip()) < 100:
                validation['warnings'].append('Document contains very little text')
                validation['recommendations'].append('Ensure the document contains substantial text content')
            elif len(text.strip()) > 50000:
                validation['warnings'].append('Document is very long')
                validation['recommendations'].append('Consider splitting very long documents')
        except Exception as e:
            validation['issues'].append(f'Could not extract text: {str(e)}')
            validation['is_valid'] = False

        return validation
